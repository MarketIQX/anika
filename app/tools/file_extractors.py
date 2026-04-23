"""Pluggable extractors for teaching-file uploads.

Supported inputs:
  .txt    plain text
  .eml    RFC-822 email (with headers + MIME bodies)
  .pdf    simple (pypdf) → complex (pdfplumber) → scanned (OCR fallback)
  .docx   Microsoft Word

Return shape (dataclass `ExtractedText`):
    text         — the extracted plain text
    metadata     — {page_count, has_tables, was_ocr, needs_ocr, source_type}

Design notes:
  - For PDFs we try pypdf first (fast). If per-page text < 50 chars we
    treat the page as likely scanned and try pdfplumber's richer extraction.
    If that's ALSO thin, we try OCR via pytesseract+pdf2image. If the
    Tesseract binary isn't installed, we DO NOT raise — we flag
    `needs_ocr=True` on the metadata so the UI can show "scanned PDF,
    install Tesseract" rather than failing the upload.
  - OCR is bounded by OCR_TIMEOUT_SECONDS (hard 60s per the spec) per page
    via a watchdog. Long multi-page scanned PDFs aren't currently chunked
    across workers — that's Phase 1B material.
"""
from __future__ import annotations

import email as email_mod
import logging
import os
from dataclasses import dataclass, field
from email import policy
from pathlib import Path
from typing import Any

import docx as _docx
import pdfplumber
import pypdf

logger = logging.getLogger(__name__)

# Per spec — anything longer per-page than this signals "real" text-layer.
PAGE_TEXT_MIN_CHARS = 50
OCR_TIMEOUT_SECONDS = 60
MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB per file


@dataclass
class ExtractedText:
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


class ExtractionError(RuntimeError):
    """Raised when the extractor cannot produce any usable text."""


class FileTooLargeError(ExtractionError):
    pass


# --------------------------------------------------------------------------
# .txt
# --------------------------------------------------------------------------

def _extract_txt(path: Path) -> ExtractedText:
    raw = path.read_bytes()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1", errors="replace")
    return ExtractedText(
        text=text,
        metadata={"source_type": "txt", "page_count": 1, "has_tables": False,
                  "was_ocr": False, "needs_ocr": False},
    )


# --------------------------------------------------------------------------
# .eml
# --------------------------------------------------------------------------

def _extract_eml(path: Path) -> ExtractedText:
    """Parse RFC-822. Collapse to:

        Subject: ...
        From:    ...
        To:      ...
        Date:    ...

        <plain-text body>
    """
    with path.open("rb") as f:
        msg = email_mod.message_from_binary_file(f, policy=policy.default)
    headers: list[str] = []
    for hdr in ("Subject", "From", "To", "Cc", "Date"):
        val = msg.get(hdr, "")
        if val:
            headers.append(f"{hdr}: {val}")
    body_parts: list[str] = []
    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type()
            if ctype == "text/plain":
                try:
                    body_parts.append(part.get_content())
                except Exception:  # noqa: BLE001
                    continue
        if not body_parts:
            for part in msg.walk():
                if part.get_content_type() == "text/html":
                    from bs4 import BeautifulSoup
                    try:
                        body_parts.append(
                            BeautifulSoup(part.get_content(), "lxml").get_text("\n").strip()
                        )
                    except Exception:  # noqa: BLE001
                        continue
    else:
        try:
            body_parts.append(msg.get_content())
        except Exception:  # noqa: BLE001
            body_parts.append(msg.as_string())
    body = "\n".join(p for p in body_parts if p)
    text = "\n".join(headers) + "\n\n" + body.strip()
    return ExtractedText(
        text=text.strip(),
        metadata={"source_type": "eml", "page_count": 1, "has_tables": False,
                  "was_ocr": False, "needs_ocr": False},
    )


# --------------------------------------------------------------------------
# .docx
# --------------------------------------------------------------------------

def _extract_docx(path: Path) -> ExtractedText:
    doc = _docx.Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    has_tables = False
    for table in doc.tables:
        has_tables = True
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return ExtractedText(
        text="\n".join(parts).strip(),
        metadata={"source_type": "docx", "page_count": 1, "has_tables": has_tables,
                  "was_ocr": False, "needs_ocr": False},
    )


# --------------------------------------------------------------------------
# .pdf — the interesting one.
# --------------------------------------------------------------------------

def _extract_pdf_simple(path: Path) -> tuple[str, int]:
    """Return (full_text, page_count) via pypdf."""
    reader = pypdf.PdfReader(str(path))
    pages = [(p.extract_text() or "") for p in reader.pages]
    return "\n\n".join(pages), len(reader.pages)


def _extract_pdf_complex(path: Path) -> tuple[str, int, bool]:
    """Return (full_text, page_count, has_tables) via pdfplumber — slower,
    handles multi-column and simple tables.
    """
    text_parts: list[str] = []
    has_tables = False
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            tables = page.extract_tables() or []
            if tables:
                has_tables = True
                for tbl in tables:
                    for row in tbl:
                        cells = [(c or "").strip() for c in row]
                        t += "\n" + " | ".join(cells)
            text_parts.append(t)
        page_count = len(pdf.pages)
    return "\n\n".join(text_parts), page_count, has_tables


def _ocr_available() -> bool:
    """Best-effort check for a usable Tesseract binary at call time."""
    try:
        import pytesseract
        # This raises if the binary isn't on PATH.
        pytesseract.get_tesseract_version()
        return True
    except Exception:  # noqa: BLE001
        return False


def _extract_pdf_ocr(path: Path, *, max_pages: int | None = None) -> tuple[str, int, bool]:
    """OCR every page as an image. Returns (text, page_count, timed_out).

    Raises ExtractionError if neither Tesseract nor Poppler are available.
    """
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError as e:
        raise ExtractionError(f"OCR library missing: {e}") from e

    images = convert_from_path(str(path))
    if max_pages is not None:
        images = images[:max_pages]

    import signal
    import threading

    timed_out = False
    # Per-page watchdog: cap the whole run at OCR_TIMEOUT_SECONDS * pages,
    # but also cap each page at OCR_TIMEOUT_SECONDS to avoid a single
    # pathological page hanging the queue. signal.SIGALRM is POSIX-only so
    # we use a threading-based watchdog for Windows compatibility.
    results: list[str] = []
    for i, img in enumerate(images):
        done: list[str] = []

        def worker(_img=img):
            try:
                done.append(pytesseract.image_to_string(_img))
            except Exception as ex:  # noqa: BLE001
                done.append(f"[OCR ERROR: {ex}]")

        t = threading.Thread(target=worker)
        t.start()
        t.join(OCR_TIMEOUT_SECONDS)
        if t.is_alive():
            timed_out = True
            results.append(f"[OCR TIMED OUT AFTER {OCR_TIMEOUT_SECONDS}s on page {i+1}]")
            # Thread will eventually finish; we move on.
            continue
        results.append(done[0] if done else "")
    return "\n\n".join(results), len(images), timed_out


def _extract_pdf(path: Path) -> ExtractedText:
    # 1. Fast pypdf pass.
    try:
        simple_text, page_count = _extract_pdf_simple(path)
    except Exception as e:  # noqa: BLE001
        raise ExtractionError(f"PDF parse failed: {e}") from e

    avg_per_page = (len(simple_text) / page_count) if page_count else 0
    if avg_per_page >= PAGE_TEXT_MIN_CHARS:
        return ExtractedText(
            text=simple_text,
            metadata={"source_type": "pdf", "page_count": page_count,
                      "has_tables": False, "was_ocr": False, "needs_ocr": False,
                      "avg_chars_per_page": int(avg_per_page)},
        )

    # 2. Richer pdfplumber pass — catches multi-column + tables.
    try:
        complex_text, page_count_c, has_tables = _extract_pdf_complex(path)
    except Exception as e:  # noqa: BLE001
        complex_text, page_count_c, has_tables = simple_text, page_count, False
        logger.warning("pdfplumber pass failed, falling back: %s", e)

    avg_per_page_c = (len(complex_text) / page_count_c) if page_count_c else 0
    if avg_per_page_c >= PAGE_TEXT_MIN_CHARS:
        return ExtractedText(
            text=complex_text,
            metadata={"source_type": "pdf", "page_count": page_count_c,
                      "has_tables": has_tables, "was_ocr": False, "needs_ocr": False,
                      "avg_chars_per_page": int(avg_per_page_c)},
        )

    # 3. Probably a scanned PDF. Try OCR if available; otherwise flag.
    if not _ocr_available():
        logger.info("Scanned PDF detected at %s but OCR not available", path.name)
        return ExtractedText(
            text=complex_text,  # the thin extraction, if any — better than empty
            metadata={"source_type": "pdf", "page_count": page_count_c,
                      "has_tables": has_tables, "was_ocr": False, "needs_ocr": True,
                      "avg_chars_per_page": int(avg_per_page_c),
                      "ocr_status": "skipped_no_tesseract"},
        )

    try:
        ocr_text, ocr_pages, timed_out = _extract_pdf_ocr(path)
    except ExtractionError as e:
        logger.warning("OCR failed for %s: %s", path.name, e)
        return ExtractedText(
            text=complex_text,
            metadata={"source_type": "pdf", "page_count": page_count_c,
                      "has_tables": has_tables, "was_ocr": False, "needs_ocr": True,
                      "ocr_status": f"error:{e}"},
        )
    return ExtractedText(
        text=ocr_text,
        metadata={"source_type": "pdf", "page_count": ocr_pages,
                  "has_tables": has_tables, "was_ocr": True, "needs_ocr": False,
                  "ocr_timed_out": timed_out},
    )


# --------------------------------------------------------------------------
# Dispatcher
# --------------------------------------------------------------------------


_DISPATCH = {
    ".txt":  _extract_txt,
    ".eml":  _extract_eml,
    ".docx": _extract_docx,
    ".pdf":  _extract_pdf,
}


def supported_extensions() -> list[str]:
    return list(_DISPATCH.keys())


def extract(path: str | Path) -> ExtractedText:
    """Extract teaching text from a file. Raises ExtractionError on failure.

    Files larger than MAX_UPLOAD_BYTES raise FileTooLargeError.
    """
    path = Path(path)
    if not path.exists():
        raise ExtractionError(f"file not found: {path}")
    size = path.stat().st_size
    if size > MAX_UPLOAD_BYTES:
        raise FileTooLargeError(
            f"file {path.name} is {size / (1024*1024):.1f} MB, exceeds "
            f"{MAX_UPLOAD_BYTES / (1024*1024):.0f} MB limit"
        )
    ext = path.suffix.lower()
    fn = _DISPATCH.get(ext)
    if fn is None:
        raise ExtractionError(
            f"unsupported extension {ext!r} — supported: {', '.join(supported_extensions())}"
        )
    return fn(path)
