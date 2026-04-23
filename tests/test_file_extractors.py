"""Tests for app/tools/file_extractors.py.

Covers the five input types we support: txt, eml, docx, simple PDF, and
scanned PDF (OCR path mocked — we don't ship Tesseract in CI)."""
from __future__ import annotations

from pathlib import Path

import pytest

from app.tools import file_extractors
from app.tools.file_extractors import ExtractedText, ExtractionError, extract


# --- Helpers --------------------------------------------------------------


def _write(tmp_path: Path, name: str, content: str | bytes) -> Path:
    p = tmp_path / name
    if isinstance(content, str):
        p.write_text(content, encoding="utf-8")
    else:
        p.write_bytes(content)
    return p


def _make_simple_pdf(tmp_path: Path, text: str) -> Path:
    """Create a single-page PDF with extractable text via pypdf's PdfWriter.

    Uses a minimal handcrafted PDF structure — pypdf can read it back with
    extract_text().
    """
    # Use pypdf's own Text embedding through PdfWriter + add_blank_page +
    # insert_page — we actually just write a valid PDF with a /Contents
    # stream containing a Tj operator. Simpler: use the `fpdf` library?
    # We don't want to pull in another dep. Instead, use pdfplumber's
    # test-friendly approach: generate via pypdf + raw content stream.
    #
    # Simplest reliable path: use reportlab-style tiny PDF via pypdf's
    # own API. pypdf doesn't create; it reads. So we write a minimal PDF
    # by hand.
    pdf_bytes = _minimal_text_pdf(text)
    return _write(tmp_path, "simple.pdf", pdf_bytes)


def _minimal_text_pdf(text: str) -> bytes:
    """Return bytes of a one-page PDF containing `text` as a text stream.

    Hand-rolled so we don't need reportlab. Valid enough for pypdf's
    extract_text() to return the string.
    """
    content = f"BT /F1 12 Tf 50 750 Td ({text}) Tj ET"
    content_bytes = content.encode("latin-1")
    stream = (
        b"<< /Length "
        + str(len(content_bytes)).encode()
        + b" >>\nstream\n"
        + content_bytes
        + b"\nendstream"
    )
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj",
        b"2 0 obj << /Type /Pages /Count 1 /Kids [3 0 R] >> endobj",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj",
        b"4 0 obj " + stream + b" endobj",
        b"5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj",
    ]
    out = b"%PDF-1.4\n"
    offsets = []
    for obj in objects:
        offsets.append(len(out))
        out += obj + b"\n"
    xref_offset = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        b"trailer << /Size 6 /Root 1 0 R >>\nstartxref\n"
        + str(xref_offset).encode()
        + b"\n%%EOF\n"
    )
    return out


# --- txt -----------------------------------------------------------------


def test_extract_txt_reads_utf8(tmp_path):
    p = _write(tmp_path, "note.txt", "Hello Anika — ñice to meet you.")
    r = extract(p)
    assert isinstance(r, ExtractedText)
    assert "Hello Anika" in r.text
    assert r.metadata["source_type"] == "txt"


def test_extract_txt_falls_back_to_latin1(tmp_path):
    # Non-UTF-8 bytes — should not raise.
    p = _write(tmp_path, "bad.txt", b"\xff\xfeHello latin-1 fallback")
    r = extract(p)
    assert "Hello" in r.text


# --- eml -----------------------------------------------------------------


def test_extract_eml_with_headers(tmp_path):
    raw = (
        "From: sender@x.com\n"
        "To: prakasha@balakrishnaandco.com\n"
        "Subject: NRI query\n"
        "Date: Mon, 1 Apr 2026 10:00:00 +0530\n"
        "Content-Type: text/plain; charset=utf-8\n"
        "\n"
        "Dear Sir,\n\nI am an NRI and need help with my ITR-2.\n\nThanks."
    )
    p = _write(tmp_path, "query.eml", raw)
    r = extract(p)
    assert "Subject: NRI query" in r.text
    assert "From: sender@x.com" in r.text
    assert "NRI and need help" in r.text
    assert r.metadata["source_type"] == "eml"


# --- docx ----------------------------------------------------------------


def test_extract_docx_reads_paragraphs_and_tables(tmp_path):
    import docx as _docx

    doc = _docx.Document()
    doc.add_paragraph("Heading line")
    doc.add_paragraph("Second paragraph with detail.")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Fee"
    t.cell(0, 1).text = "Amount"
    t.cell(1, 0).text = "ITR filing"
    t.cell(1, 1).text = "₹8,000"
    p = tmp_path / "teach.docx"
    doc.save(str(p))

    r = extract(p)
    assert "Heading line" in r.text
    assert "Second paragraph" in r.text
    assert "ITR filing" in r.text
    assert "₹8,000" in r.text
    assert r.metadata["has_tables"] is True


# --- pdf (simple) --------------------------------------------------------


def test_extract_simple_pdf(tmp_path):
    p = _make_simple_pdf(tmp_path, "Anika teaching content line 1. Phrase-level voice sample.")
    r = extract(p)
    assert "Anika" in r.text
    assert r.metadata["source_type"] == "pdf"
    assert r.metadata["was_ocr"] is False
    # 50+ chars, so OCR is skipped.
    assert r.metadata["needs_ocr"] is False


# --- pdf (scanned → OCR path, with OCR unavailable) ---------------------


def test_extract_scanned_pdf_flags_needs_ocr_when_no_tesseract(tmp_path, monkeypatch):
    """If the text layer is thin AND Tesseract isn't available, the extractor
    must flag `needs_ocr=True` instead of raising."""
    # A "scanned" PDF in our heuristic = < 50 chars per page of extractable
    # text. Use a near-empty text stream.
    p = _make_simple_pdf(tmp_path, "x")  # 1-char text layer
    monkeypatch.setattr(file_extractors, "_ocr_available", lambda: False)
    r = extract(p)
    assert r.metadata["was_ocr"] is False
    assert r.metadata["needs_ocr"] is True
    assert "skipped_no_tesseract" in (r.metadata.get("ocr_status") or "")


def test_extract_scanned_pdf_ocr_path_with_mock(tmp_path, monkeypatch):
    """When Tesseract IS available, the extractor calls the OCR code path
    and uses its output."""
    p = _make_simple_pdf(tmp_path, "y")  # thin text → OCR trigger

    monkeypatch.setattr(file_extractors, "_ocr_available", lambda: True)

    def fake_ocr(path, *, max_pages=None):
        return "OCR-derived text layer from mock.", 1, False

    monkeypatch.setattr(file_extractors, "_extract_pdf_ocr", fake_ocr)

    r = extract(p)
    assert "OCR-derived text" in r.text
    assert r.metadata["was_ocr"] is True
    assert r.metadata["needs_ocr"] is False


# --- size limit ----------------------------------------------------------


def test_file_too_large(tmp_path, monkeypatch):
    # Stub the limit down so we don't need a real 50 MB file.
    monkeypatch.setattr(file_extractors, "MAX_UPLOAD_BYTES", 10)
    p = _write(tmp_path, "big.txt", "x" * 1024)
    with pytest.raises(file_extractors.FileTooLargeError):
        extract(p)


# --- unknown extension --------------------------------------------------


def test_unknown_extension_raises(tmp_path):
    p = _write(tmp_path, "unknown.xyz", "content")
    with pytest.raises(ExtractionError):
        extract(p)
